#!/usr/bin/env python3
"""
Convert Markdown files to DOCX format.

Usage:
    python3 md_to_docx.py --input data/output/letters/
    python3 md_to_docx.py --file data/output/letters/consent-acme.md
    python3 md_to_docx.py --input data/output/reports/

Outputs .docx files alongside the .md files.
"""

import argparse
import re
import sys
from pathlib import Path

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH


def parse_md_line(line: str) -> dict:
    """Classify a markdown line by type."""
    stripped = line.strip()
    if stripped.startswith('## '):
        return {'type': 'heading2', 'text': stripped[3:]}
    if stripped.startswith('# '):
        return {'type': 'heading1', 'text': stripped[2:]}
    if stripped.startswith('> '):
        return {'type': 'quote', 'text': stripped[2:]}
    if re.match(r'^\d+\.', stripped):
        return {'type': 'list_item', 'text': stripped}
    if stripped.startswith('- '):
        return {'type': 'bullet', 'text': stripped[2:]}
    if stripped.startswith('**') and stripped.endswith('**'):
        return {'type': 'bold_line', 'text': stripped[2:-2]}
    if stripped == '':
        return {'type': 'blank', 'text': ''}
    return {'type': 'paragraph', 'text': stripped}


def convert_md_to_docx(md_path: Path, docx_path: Path) -> None:
    """Convert a single markdown file to docx."""
    doc = Document()

    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    # Set margins
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.2)
        section.right_margin = Inches(1.2)

    md_text = md_path.read_text(encoding='utf-8')
    lines = md_text.split('\n')

    i = 0
    while i < len(lines):
        parsed = parse_md_line(lines[i])

        if parsed['type'] == 'blank':
            i += 1
            continue

        if parsed['type'] == 'heading1':
            p = doc.add_heading(parsed['text'], level=1)
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            i += 1
            continue

        if parsed['type'] == 'heading2':
            p = doc.add_heading(parsed['text'], level=2)
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            i += 1
            continue

        if parsed['type'] == 'quote':
            # Collect multi-line quotes
            quote_lines = [parsed['text']]
            i += 1
            while i < len(lines) and lines[i].strip().startswith('> '):
                quote_lines.append(lines[i].strip()[2:])
                i += 1
            quote_text = ' '.join(quote_lines)
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.5)
            p.paragraph_format.right_indent = Inches(0.5)
            run = p.add_run(f'"{quote_text}"')
            run.italic = True
            run.font.size = Pt(10)
            continue

        if parsed['type'] == 'list_item':
            p = doc.add_paragraph(parsed['text'], style='List Number')
            i += 1
            continue

        if parsed['type'] == 'bullet':
            p = doc.add_paragraph(parsed['text'], style='List Bullet')
            i += 1
            continue

        if parsed['type'] == 'bold_line':
            p = doc.add_paragraph()
            run = p.add_run(parsed['text'])
            run.bold = True
            i += 1
            continue

        # Regular paragraph — handle inline bold with **text**
        text = parsed['text']
        p = doc.add_paragraph()
        parts = re.split(r'(\*\*[^*]+\*\*)', text)
        for part in parts:
            if part.startswith('**') and part.endswith('**'):
                run = p.add_run(part[2:-2])
                run.bold = True
            else:
                p.add_run(part)
        i += 1

    doc.save(str(docx_path))


def process_directory(input_dir: Path) -> int:
    """Convert all .md files in a directory to .docx."""
    count = 0
    for md_file in sorted(input_dir.glob('*.md')):
        docx_file = md_file.with_suffix('.docx')
        print(f"Converting: {md_file.name} -> {docx_file.name}")
        convert_md_to_docx(md_file, docx_file)
        count += 1
    return count


def main():
    parser = argparse.ArgumentParser(description="Convert MD to DOCX")
    parser.add_argument("--input", help="Directory of .md files")
    parser.add_argument("--file", help="Single .md file to convert")
    args = parser.parse_args()

    if args.file:
        md_path = Path(args.file)
        if not md_path.exists():
            print(f"File not found: {md_path}")
            sys.exit(1)
        docx_path = md_path.with_suffix('.docx')
        convert_md_to_docx(md_path, docx_path)
        print(f"Converted: {md_path.name} -> {docx_path.name}")

    elif args.input:
        input_dir = Path(args.input)
        if not input_dir.is_dir():
            print(f"Directory not found: {input_dir}")
            sys.exit(1)
        count = process_directory(input_dir)
        print(f"\nConverted {count} files")

    else:
        parser.print_help()


if __name__ == "__main__":
    main()
